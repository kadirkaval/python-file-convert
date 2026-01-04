from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document(data_json):
    # Belge oluştur
    doc = Document()

    # --- Başlık Kısmı ---
    title = doc.add_heading(data_json["baslik"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Alt Başlık
    subtitle = doc.add_paragraph(f"Tarih: {data_json['tarih']} | Ders: {data_json['ders']}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER 

    # --- Sorular ve Cevaplar Döngüsü ---
    for bolum in data_json["bolumler"]:
        # Bölüm Başlığı
        doc.add_heading(bolum["bolum_adi"], level=1)
        
        # O bölümdeki maddeler
        for item in bolum["icerik"]:
            p = doc.add_paragraph()
            # Soru kısmı kalın
            runner = p.add_run(f"{item['soru']}: ")
            runner.bold = True
            # Cevap kısmı normal
            p.add_run(f"{item['cevap']}")

    # --- Kaydetme ---
    dosya_adi = f"{data_json['dosya_adi']}.docx"
    doc.save(dosya_adi)
    print(f"✅ Belge başarıyla oluşturuldu: {dosya_adi}")

# --- VERİ SETİ ---
# Burayı değiştirebilirsin
veri_seti = {}

# --- ÇALIŞTIRMA KISMI ---
if __name__ == "__main__":
    # Fonksiyona doğrudan değişkeni gönderiyoruz, "veri_seti=" demeden.
    create_document(veri_seti)